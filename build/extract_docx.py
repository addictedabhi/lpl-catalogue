from __future__ import annotations
import os
import docx
from docx.oxml.ns import qn
from build.parse_text import parse_product_lines

# (category_name, category_id) in display order. Index = position in docx.tables.
CATEGORY_TABLES = [
    ("Rakhi", "rakhi"),
    ("Bathrobes", "bathrobes"),
    ("Brooches", "brooches"),
    ("Hair Clips", "hairclips"),
    ("Collars", "collars"),
]
# docx table indices for the five category tables (verified against v3 docx).
_TABLE_INDEX = {"rakhi": 3, "bathrobes": 4, "brooches": 5, "hairclips": 6, "collars": 7}


def _cell_lines(cell) -> list[str]:
    return [p.text for p in cell.paragraphs]


def _cell_image_blob(cell, part):
    """Return (blob, ext) for the first embedded image in a table cell, or None."""
    for blip in cell._element.findall(".//" + qn("a:blip")):
        rid = blip.get(qn("r:embed"))
        if rid and rid in part.rels:
            img = part.rels[rid].target_part
            ext = os.path.splitext(img.partname)[1] or ".png"
            return img.blob, ext
    return None


def extract_catalogue(docx_path: str, image_out_dir: str) -> list[dict]:
    doc = docx.Document(docx_path)
    part = doc.part
    cats: list[dict] = []

    for name, cid in CATEGORY_TABLES:
        table = doc.tables[_TABLE_INDEX[cid]]
        cat_dir = os.path.join(image_out_dir, cid)
        os.makedirs(cat_dir, exist_ok=True)

        # Row 0: [title, blurb]
        header_cells = table.rows[0].cells
        blurb = header_cells[-1].text.strip() if len(header_cells) > 1 else ""

        products: list[dict] = []
        idx = 0
        seen_rows = set()
        for row in table.rows[1:]:
            row_key = id(row._tr)
            if row_key in seen_rows:
                continue
            seen_rows.add(row_key)

            # De-duplicate merged cells within the row
            cells = []
            seen_cell_ids = set()
            for cell in row.cells:
                cid_key = id(cell._element)
                if cid_key not in seen_cell_ids:
                    seen_cell_ids.add(cid_key)
                    cells.append(cell)

            # Identify text cell and image cell
            text_cell = None
            img_cell = None
            for cell in cells:
                if cell._element.findall(".//" + qn("a:blip")):
                    img_cell = cell
                else:
                    text_cell = cell

            if text_cell is None:
                continue
            rec = parse_product_lines(_cell_lines(text_cell))
            if not rec:
                continue

            docx_image = None
            if img_cell:
                img = _cell_image_blob(img_cell, part)
                if img:
                    blob, ext = img
                    path = os.path.join(cat_dir, f"{idx}{ext}")
                    with open(path, "wb") as f:
                        f.write(blob)
                    docx_image = path.replace("\\", "/")

            rec["category"] = name
            rec["docx_image"] = docx_image
            products.append(rec)
            idx += 1

        cats.append({"name": name, "id": cid, "blurb": blurb, "products": products})
    return cats
